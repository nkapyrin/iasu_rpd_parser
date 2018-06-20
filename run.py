#encoding:utf-8
import os
from os.path import isfile, join
import docx
import codecs
from shorten_bib_entry import shorten

from get_up import get_up

bib = set()
dict_by_discipline = {}

up_field_names = [u"﻿Наименование",u"Профиль/Специализация",u"Шифр профиля/специализации",u"Площадка",u"Год",\
                  u"Форма обучения",u"Уровень обучения",u"Для иностранных групп",u"Выпускающая кафедра",\
                  u"Статус разработки",u"Утверждено проректором",u"Приведен к шаблону",u"Скан прикреплен",\
                  u"Дата отправки на утверждение",u"Дата утверждения",u"Код",u"Нет набора",u"Для лицензирования" ]
rpd_field_names = [ u"Наименование", u"Дисциплина", u"Учебный план", u"Шифр", u"Выпускающая кафедра",\
                    u"Обеспечивающая кафедра",u"Код", u"Год поступления", u"Статус разработки", \
                    u"Дата отправки на утверждение",u"Дата утверждения",u"Утверждено проректором"]
rpd_by_plan = {}

up_info = {}
with codecs.open( "up_list.txt", 'r', encoding="utf-8" ) as upf:
    up_lines = [  l.replace('\n','').replace('\r','').strip().strip(';') for l in upf.readlines() ]
    up_lst = [ l.split(';') for l in up_lines[1:] ]
    up_srt = sorted( up_lst, key=lambda e:(e[1],e[4]), reverse=True )

    for up_fields in up_srt: #in up_lines[1:]:
        #print len(up_line.split(';'))
        up_tmp = dict( zip(up_field_names, [s.replace('\n','').replace('\r','').strip() for s in up_fields] )  )
        file_nb = str(int( up_tmp[u"Код"] ))
        # Пройтись по всем РПД этого УП
        if not isfile( join(u'plan',file_nb+".csv") ):
           print join(u'plan',file_nb+".csv"), "missing"
           continue;
        else: print join(u'plan',file_nb+".csv"), "OK"

        #up_id =  up[u"Шифр профиля/специализации"] + u" " + up[u"Год"].replace('/','-')
        up_id =  up_tmp[u"Шифр профиля/специализации"]
        if up_id not in rpd_by_plan.keys():
            rpd_by_plan[ up_id ] = up_tmp
            rpd_by_plan[ up_id ]["rpd"] = {}
            rpd_by_plan[ up_id ]["id"] = up_id
            rpd_by_plan[ up_id ]["header"] = u"Специализация " + up_tmp[u"Шифр профиля/специализации"] + " " + up_tmp[u"Профиль/Специализация"] # " " + up_tmp[u"Год"]

        if up_id not in rpd_by_plan.keys(): rpd_by_plan[up_id] = {}
        for folder in [ 'lists_disciplines', 'lists_internship' ]:

          filename = ''
          if isfile( join( folder, file_nb+".csv" )): filename = file_nb+".csv"
          elif isfile( join( folder, file_nb+".txt" )): filename = file_nb+".txt"
          else:
            print "No RPD" + file_nb 
            continue

          with codecs.open( join( folder, filename), 'r', encoding="utf-8" ) as f:
            if '.csv' in filename: lines = [  l.replace('\n','').replace('\r','').strip().strip(';') for l in f.readlines() ]
            else: lines = [  l.replace('\n','').replace('\r','').strip().strip('\t') for l in f.readlines() ]
            for line in lines[1:]:
                if '.csv' in filename: rpd_tmp = dict( zip(rpd_field_names, line.split(';'))  )
                else: rpd_tmp = dict( zip(rpd_field_names, line.split('\t'))  )
                #if u"305" not in rpd[u"Выпускающая кафедра"]: continue    # если нужно исключить РПД других кафедр 
                rpd_tmp[u"Основная литература"] = []
                rpd_tmp[u"Дополнительная литература"] = []
                
                fnRPD = u"rpd"+rpd_tmp[u"Код"]+".txt"
                idRPD = rpd_tmp[u"Дисциплина"]

                if idRPD not in rpd_by_plan[ up_id ]["rpd"].keys(): rpd_by_plan[ up_id ]["rpd"][ idRPD ] = rpd_tmp
                else: continue # Не будем обрабатывать более старые РПД (УП отсортированы по годам)

                # Если такого файла нет, ничего не поделаешь
                if not isfile( join(u'rpd_txt_files',fnRPD) ):
                    print '   ', join(u'rpd_txt_files',fnRPD), "missing"
                    continue;
                # Прочитать файл РПД и найти всю нужную информацию
                with codecs.open( join(u'rpd_txt_files',fnRPD),'r', encoding="utf-8" ) as rpdf:
                    rpd_lines = [  l.replace('\n','').replace('\r','').replace(u'¬','').strip().strip(';') for l in rpdf.readlines() ]
                    cur_discip_nb = ""; cur_discip_name = "";
                    writing_main_bibliography = 0; writing_extra_bibliography = 0; writing_ecatalog_bib = 0; #next_is_discipline_name = 0;
                    # Просмотреть все строчки файла РПД
                    for rpdl in rpd_lines:
                        if len(rpdl.strip()) == 0: continue
                        elif u"а)основная литература:" in rpdl:
                          writing_extra_bibliography = 0; writing_main_bibliography = 1;
                        elif u"б)дополнительная литература:" in rpdl:
                          writing_extra_bibliography = 1; writing_main_bibliography = 0;
                        elif u"Литература из электронного каталога" in rpdl or u"Дополнительная литература" in rpdl: continue;
                        elif l.strip(' ') == "" or u"ПЕРЕЧЕНЬ РЕСУРСОВ ИНФОРМАЦИОННО" in rpdl or u'МАТЕРИАЛЬНО-ТЕХНИЧЕСКОЕ ОБЕСПЕЧЕНИЕ ПРАКТИКИ' in rpdl or u"Интернет-ресурсы" in rpdl or u"Интернет-сайты" in rpdl or u"Периодические издания" in rpdl:
                          writing_extra_bibliography = 0; writing_main_bibliography = 0;
                        elif writing_main_bibliography == 1:
                            txt = rpdl.strip()
                            if txt[0] in u'1234567890':
                              separator = u'.' if txt.find(u'.') < 4 else u')'
                              txt = txt[ txt.find(separator)+1: ].strip()
                            if len(txt.strip()) > 0: rpd_by_plan[ up_id ]["rpd"][idRPD][u"Основная литература"].append( txt )
                        elif writing_extra_bibliography==1:
                            txt = rpdl.strip()
                            if txt[0] in u'1234567890':
                              separator = u'.' if txt.find(u'.') < 4 else u')'
                              txt = txt[ txt.find(separator)+1: ].strip()
                              #txt = txt[ txt.find(u'.')+1: ].strip()
                            if len(txt.strip()) > 0: rpd_by_plan[ up_id ]["rpd"][idRPD][u"Дополнительная литература"].append( txt )

# Сохранить на диск
with codecs.open( "books.txt", 'w' ) as f:
  for up_id in rpd_by_plan.keys():
    #f.write( (u"\n" + up_id + u"\n").encode('utf-8') )
    for rpd_id in rpd_by_plan[up_id]['rpd']:
      #f.write( (rpd_id + '\n').encode('utf-8') )
      rpd = rpd_by_plan[up_id]['rpd'][rpd_id]
      for b in (rpd[u"Основная литература"]+rpd[u"Дополнительная литература"]):
          f.write( (b + u"\n").encode('utf-8') )

# Сохранить на диск
with codecs.open( "records.txt", 'w' ) as f:
  for up_id in rpd_by_plan.keys():
    for rpd_id in rpd_by_plan[up_id]['rpd']:
      rpd = rpd_by_plan[up_id]['rpd'][rpd_id]
      #f.write( (rpd_id + " " + rpd[u"Код"] + '\n').encode('utf-8') )
      for b in (rpd[u"Основная литература"] + rpd[u"Дополнительная литература"]):
          if len( shorten(b) ) > 0: f.write( (shorten(b) + '\n').encode('utf-8') )

# Загрузить данные из библиотеки
os.system( 'python query_library.py' )


#print '---'
#for up_id in rpd_by_plan.keys(): print up_id
#print '---'


books_info = {}
keys_lst = [ u"Код", u"Тип", u"Запись", u"Авторы", u"Экземпляры", u"Шифры", u"Ссылка" ]
with codecs.open( "books_records_out.txt", 'r' ) as f:
    lines = [  l.replace('\n','').replace('\r','').decode("utf-8") for l in f.readlines() ]
    lines_lst = [ l.split('|') for l in lines ]
    for fields in lines_lst:
        books_info[ fields[0] ] = dict( zip(keys_lst, fields))
        ex = books_info[ fields[0] ][u"Экземпляры"]
        books_info[ fields[0] ][u"Всего"] = ex[ ex.find(u"Всего")+7:ex.find(u",") ]


list_books_in_lib = []
list_books_not_in_lib = []

for up_id in rpd_by_plan.keys():
    up = rpd_by_plan[up_id]
    file_nb = str(int( up[u"Код"] ))
    template = docx.Document('report_template.docx');
    print up_id
    for tb in template.tables:
        if tb.rows[0].cells[0].paragraphs[0].text == u"#TABLE_BY_PLAN_INSERT_PLAN_ID_AND_NAME#":
            tb.rows[0].cells[0].paragraphs[0].text = up["header"]
            src_row = tb.rows[-1]; src_cells = src_row.cells;
            for ii,idRPD in enumerate( sorted(rpd_by_plan[ up_id ]["rpd"].keys()) ):
                print idRPD
                rpd = rpd_by_plan[ up_id ]["rpd"][idRPD]
                first_row_cells, last_row_cells = None, None
                for ib,b in enumerate(rpd[u"Основная литература"] + rpd[u"Дополнительная литература"]):
                  tb.add_row()
                  #print '>>>', b
                  if ib==0: first_row_cells = tb.rows[-1].cells
                  last_row_cells = tb.rows[-1].cells
                  
                  # Занести библиографическую запись в третью колонку шаблонного документа
                  last_row_cells[3].text = b

                  # Если текущая запись есть в перечне, подставим количество книг
                  if short in books_info.keys():
                      last_row_cells[4].text = books_info[short][u"Всего"]
                      if len( books_info[short][u"Ссылка"] ) > 0: last_row_cells[5].text = u"✓"
                      list_books_in_lib.append( short )
                  # Если такой записи нет, попробуем подставить запись, где книга издана в другом году
                  elif '$' in short and short.split('$')[0] in books_info.keys():
                      short = short.split('$')[0]
                      last_row_cells[4].text = books_info[short][u"Всего"] + u'*'
                      if len( books_info[short][u"Ссылка"] ) > 0: last_row_cells[5].text = u"✓"
                      list_books_in_lib.append( short )
                  else: list_books_not_in_lib.append( short )
                  #
                  for iic,c in enumerate(tb.rows[-1].cells): c.paragraphs[0].style = src_cells[iic].paragraphs[0].style

                if first_row_cells and last_row_cells:
                  A = first_row_cells[0].merge( last_row_cells[0] ); A.text = str(ii+1)
                  B = first_row_cells[2].merge( last_row_cells[2] ); B.text = rpd[ u"Дисциплина" ] + "\n" + rpd[u"Код"]
                  first_row_cells[1].merge( last_row_cells[1] );
                  first_row_cells[6].merge( last_row_cells[6] );

            print '\n'
            tb._tbl.remove(src_row._tr)
    template.save( u"out/Литература по предметам %s.docx" % up_id )

print "missing books (%d, %d unique):" % ( len(list_books_not_in_lib), len( list(set(list_books_not_in_lib))) )
print "OK books (%d):" % len(list_books_in_lib)
